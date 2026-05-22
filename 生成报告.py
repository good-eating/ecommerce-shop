# -*- coding: utf-8 -*-
"""生成课程设计报告"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# ========== 配置信息 - 请修改为你自己的 ==========
STUDENT_NAME = "你的姓名"
STUDENT_ID = "你的学号"
TEACHER = "布社辉"
GIT_URL = "你的GitHub仓库地址"
DEPLOY_URL = "http://39.96.40.38"
TEST_ACCOUNT = "admin / admin123"
# =============================================

doc = Document("D:/study/junior2/shop/课程设计报告模板.docx")

def add_para(doc, text, bold=False, size=None, align=None, font_name=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if align: p.alignment = align
    if font_name: run.font.name = font_name
    return p

def find_and_replace_table(doc, replacements):
    """在表格中查找替换文本"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run(str(value))
                            run.font.size = Pt(10.5)

# 替换基本信息
find_and_replace_table(doc, {
    "你的姓名": STUDENT_NAME,
    "你的学号": STUDENT_ID,
    "布\n\n社辉": TEACHER,
    "布\n \n社辉": TEACHER,
    "你的GitHub仓库地址": GIT_URL,
    "http://39.96.40.38": DEPLOY_URL,
    "admin / admin123": TEST_ACCOUNT,
    "1\n.0": "1.0",
})

# ========== 需求分析报告 ==========
# 找到"概述"段落，在其后插入内容
for para in doc.paragraphs:
    if "概述" in para.text and "【基本要求】" not in para.text:
        # 在概述后面添加内容
        pass

# ========== 系统设计报告 ==========

# ========== 系统实现报告 ==========

# ========== AI工具使用记录 ==========

# ========== 测试报告 ==========

output_path = f"D:/study/junior2/shop/{STUDENT_ID}_{STUDENT_NAME}.docx"
doc.save(output_path)
print(f"报告已生成: {output_path}")
